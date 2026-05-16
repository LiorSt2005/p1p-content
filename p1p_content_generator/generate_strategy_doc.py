from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x19, 0x29)
BLUE   = RGBColor(0x2A, 0x6D, 0xD9)
CYAN   = RGBColor(0x06, 0xB6, 0xD4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x44, 0x44, 0x44)
LGREY  = RGBColor(0xF2, 0xF4, 0xF8)
RED    = RGBColor(0xC0, 0x39, 0x2B)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: shade table cell ──────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ──────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{ side}')
            tag.set(qn('w:val'),   kwargs[side].get('val',   'single'))
            tag.set(qn('w:sz'),    kwargs[side].get('sz',    '4'))
            tag.set(qn('w:space'), kwargs[side].get('space', '0'))
            tag.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: remove table borders ─────────────────────────────────────────────
def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

# ── Helper: heading styles ────────────────────────────────────────────────────
def add_h1(text):
    p = doc.add_paragraph()
    shade_paragraph(p, '0D1929')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = WHITE
    return p

def add_h2(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    return p

def add_body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size   = Pt(10.5)
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = GREY
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GREY
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = GREY

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2A6DD9')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_table(headers, rows, header_color='0D1929', alt_color='F2F4F8'):
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = alt_color if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_callout(text, bg='EBF4FF', border_color='2A6DD9'):
    p = doc.add_paragraph()
    shade_paragraph(p, bg)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    # left border via pPr pBdr
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '20')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY
    run.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
shade_paragraph(p, '0D1929')
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('P1P — Estratégia de Publicação LinkedIn & Instagram')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = WHITE

p2 = doc.add_paragraph()
shade_paragraph(p2, '0D1929')
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(30)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Guia completo — calendário, horários, impulsionamento e amplificação')
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x5B, 0x8D, 0xEF)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# 1. AUDIENCE
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  1. Perfil do Público no LinkedIn')
doc.add_paragraph()

add_h3('Alvos primários')
add_bullet('CEO / Sócio-Diretor de empresas com 30–300 funcionários')
add_bullet('CFO / Diretor Financeiro / Controller')
add_bullet('COO / Gerente Administrativo-Financeiro')

add_h3('Comportamento no LinkedIn (Brasil)')
add_bullet('Verificam o LinkedIn no celular pela manhã antes da primeira reunião')
add_bullet('Leem durante o intervalo do almoço')
add_bullet('Raramente engajam em fins de semana ou após as 19h')
add_bullet('Respondem a conteúdo baseado em dor — sentem os problemas que descrevemos')
add_bullet('Compartilham conteúdo que os faz parecer inteligentes para sua rede')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. POSTING SCHEDULE
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  2. Calendário de Publicações — Dias e Horários (BRT)')
doc.add_paragraph()

add_callout('Publique 3x por semana. Mais do que isso e o engajamento por post cai. Menos e o algoritmo esquece você.')
doc.add_paragraph()

add_table(
    ['Dia', 'Horário (BRT)', 'Formato', 'Plataforma', 'Track'],
    [
        ['Terça-feira', '08:00', 'Post LinkedIn (texto longo)', 'LinkedIn', 'BPO ou BI'],
        ['Quarta-feira', '12:00 (LI) / 19:00 (IG)', 'Infográfico (imagem)', 'LinkedIn + Instagram', 'BI ou BPO'],
        ['Quinta-feira', '08:30 (LI) / 18:00 (IG)', 'Carrossel', 'LinkedIn + Instagram', 'BPO ou BI'],
    ]
)

add_body('O Instagram atinge o pico mais tarde do que o LinkedIn — seu público verifica o IG à noite e o LinkedIn antes do trabalho. Publique o mesmo conteúdo duas vezes: de manhã no LinkedIn, programe para o Instagram à tarde.')

add_h3('Alternância de tracks por semana')
add_bullet('Semana A (semanas ímpares): foco BPO — 2 posts BPO, 1 BI')
add_bullet('Semana B (semanas pares): foco BI — 2 posts BI, 1 BPO')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. CONTENT ROTATION
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  3. Rotação de Conteúdo — Como Usar os 6 Arquivos por Semana')
doc.add_paragraph()

add_body('Você tem por semana temática: txt (post LinkedIn) + infográfico (PNG) + carrossel (ZIP com slides). Não publique tudo na mesma semana. Distribua assim:')

add_table(
    ['Semana', 'O que publicar'],
    [
        ['Semana 1', 'Post texto LinkedIn + Infográfico'],
        ['Semana 2', 'Carrossel + Infográfico reciclado com nova legenda'],
    ]
)

add_body('Isso dobra a vida de cada peça. Um carrossel publicado 10 dias após o post de texto atinge um público parcialmente diferente devido à variação do algoritmo do LinkedIn.')

add_h3('Exemplo de mês com o conteúdo da Semana 3')
add_table(
    ['Data', 'Post', 'Formato'],
    [
        ['Ter 19/mai', 'BPO texto — Governança', 'Post texto longo'],
        ['Qua 21/mai', 'BI infográfico — Dados', 'Imagem PNG'],
        ['Qui 22/mai', 'BPO carrossel — Governança', 'Carrossel'],
        ['Ter 26/mai', 'BI texto — Integração', 'Post texto longo'],
        ['Qui 28/mai', 'BPO infográfico — Governança', 'PNG (nova legenda)'],
    ]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ORGANIC REACH
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  4. Como Publicar para Máximo Alcance Orgânico')
doc.add_paragraph()

add_callout('Sempre publique nativamente — nunca agende via ferramentas de terceiros (Buffer, Hootsuite, etc.) para os primeiros 3 dias. O LinkedIn penaliza metadados de agendadores externos. Use o agendador nativo do LinkedIn (sem penalidade).')
doc.add_paragraph()

add_h3('Os primeiros 60 minutos são tudo')
add_body('O LinkedIn mede a velocidade do engajamento inicial. Após publicar:')
add_bullet('Responda a todos os comentários em até 1 hora')
add_bullet('Peça a um colega para curtir e comentar imediatamente (até uma pessoa ajuda o sinal)')
add_bullet('Comente no seu próprio post 5 minutos após a publicação com uma visão adicional ou uma pergunta — isso renotifica seus seguidores')

add_h3('Hashtags — LinkedIn')
add_body('Use exatamente 3–5 hashtags. Adicione no corpo do post.')

add_table(
    ['Track', 'Hashtags recomendadas'],
    [
        ['BPO', '#BPOFinanceiro #GovernançaFinanceira #GestãoFinanceira'],
        ['BI',  '#BusinessIntelligence #PowerBI #IntegracaoDeDados'],
    ]
)

add_callout('Não adicione links no corpo do post. Se precisar linkar, coloque no primeiro comentário. O LinkedIn suprime posts com links externos no corpo.')

add_h3('Hashtags — Instagram (use 8–12)')
add_table(
    ['Track', 'Hashtags recomendadas'],
    [
        ['BPO', '#BPOFinanceiro #GovernançaFinanceira #GestãoFinanceira #Nibo #ControleFinanceiro #CFO #EmpresasBrasileiras #PequenasMediasEmpresas #Financeiro #Empreendedorismo'],
        ['BI',  '#BusinessIntelligence #PowerBI #IntegracaoDeDados #Dashboards #DataDriven #GestãoEmpresarial #TransformacaoDigital #Dados #EmpresasBrasileiras'],
    ]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CAROUSEL CROSS-PLATFORM
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  5. Carrossel: Distribuição em Duas Plataformas')
doc.add_paragraph()

add_body('Os arquivos de carrossel já exportam em 1080×1080px PNG (um arquivo por slide), que é exatamente o formato nativo do Instagram. Os mesmos slides vão para as duas plataformas sem nenhum trabalho extra.')

add_table(
    ['Plataforma', 'Como subir', 'Notas'],
    [
        ['Instagram', 'Faça upload dos slides diretamente como post carrossel (até 10 imagens)', 'Formato nativo, sem conversão'],
        ['LinkedIn',  'Opção A: post multi-imagem (selecione todos os PNGs em ordem)\nOpção B: combine em PDF via Canva para o carrossel de documento clássico', 'PDF gera mais alcance no LinkedIn'],
    ]
)

add_h3('Legenda do Instagram (diferente do LinkedIn)')
add_body('O texto longo do LinkedIn é longo demais para o Instagram. Escreva uma legenda curta para o IG:')
add_bullet('Linha 1: gancho (mesmo texto do slide 1)')
add_bullet('Linhas 2–3: expansão em 2 frases')
add_bullet('Linha 4: CTA — "Deslize para ver →"')
add_bullet('Linha 5: "Link na bio para conversar com um especialista"')

add_h3('Bio do Instagram')
add_callout('BPO Financeiro + BI Estratégico para médias empresas 🇧🇷\nParceiros oficiais @nibo.app\n👇 Fale com um especialista')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. BOOSTING — LINKEDIN
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  6. Impulsionamento — LinkedIn Campaign Manager')
doc.add_paragraph()

add_callout('O que impulsionar: apenas posts que já performaram bem organicamente (100+ impressões, 5+ reações). Nunca impulsione um post morto.')
doc.add_paragraph()

add_h3('Como configurar')
add_bullet('Acesse o LinkedIn Campaign Manager')
add_bullet('Crie uma campanha → Objetivo: Lead Generation (não Brand Awareness)')
add_bullet('Formato: Sponsored Content (impulsiona seu post existente)')
add_bullet('Adicione um Lead Gen Form — pré-preenchido com dados do perfil do LinkedIn')

add_h3('Texto do Lead Gen Form sugerido')
add_callout('Diagnóstico gratuito para o seu financeiro\nCampos: Nome / E-mail / Empresa / Cargo\nBotão CTA: "Falar com especialista"')
doc.add_paragraph()

add_h3('Configuração de segmentação (salve como template)')
add_table(
    ['Parâmetro', 'Configuração'],
    [
        ['Localização',     'Brasil'],
        ['Tamanho da empresa', '51–500 funcionários'],
        ['Cargos',          'CEO, CFO, Diretor Financeiro, Controller, Gerente Financeiro, Sócio, Proprietário'],
        ['Setores',         'Serviços, Indústria, Varejo, Construção Civil, Tecnologia'],
        ['Excluir',         'Governo, Ensino, ONG'],
        ['Idioma',          'Português'],
    ]
)

add_h3('Orçamento recomendado — LinkedIn')
add_table(
    ['Fase', 'Orçamento/dia', 'Duração', 'Objetivo'],
    [
        ['Teste (Mês 1)',    'R$30–50/dia',   '7 dias por post',  'Aprender o que converte'],
        ['Escala (Mês 2+)', 'R$80–120/dia',  '14 dias',          'Fluxo consistente de leads'],
    ]
)

add_h3('Qual conteúdo impulsionar')
add_bullet('Carrossel BPO → lead gen form → "Diagnóstico financeiro gratuito"')
add_bullet('Infográfico BI → lead gen form → "Avaliação do seu projeto de BI"')
add_body('Carrosséis superam imagens únicas para lead gen porque o tempo de permanência sinaliza qualidade para o algoritmo, o que então reduz seu custo por clique.')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. BOOSTING — INSTAGRAM / META
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  7. Impulsionamento — Meta Ads Manager (Instagram)')
doc.add_paragraph()

add_table(
    ['Plataforma', 'Objetivo', 'Conteúdo', 'Intenção do público'],
    [
        ['LinkedIn',   'Lead Gen Form',    'Infográfico / carrossel', 'Alta — pensando ativamente em negócios'],
        ['Instagram',  'Tráfego / Awareness', 'Carrossel',            'Média — aquece, retargeta depois'],
    ]
)

add_h3('Retargeting no Instagram')
add_body('Qualquer pessoa que deslizar por 3+ slides do seu carrossel é um lead quente. No Meta Ads Manager, crie um Público Personalizado → Engajou com post do Instagram e retargete-os com um anúncio de CTA direto 7 dias depois. Este é um uso muito eficiente do conteúdo de carrossel que você já está produzindo.')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 8. AMPLIFICATION
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  8. Amplificação — Além das Publicações')
doc.add_paragraph()

add_h3('WhatsApp Business')
add_bullet('Após cada post do LinkedIn, compartilhe o PNG do infográfico diretamente na sua lista de broadcast do WhatsApp Business')
add_bullet('Legenda: teaser de 2 linhas + "Conteúdo completo no LinkedIn" (direciona tráfego para o perfil)')
add_bullet('Isso não custa nada e reutiliza o que já construímos')

add_h3('Advocacy da equipe')
add_bullet('Peça a cada membro da equipe P1P para compartilhar o post (não apenas curtir — compartilhar com comentário)')
add_bullet('O LinkedIn multiplica o alcance em ~6x quando compartilhado com comentário vs. compartilhamentos silenciosos')

add_h3('Tag @Nibo')
add_bullet('Ao publicar conteúdo BPO, marque @Nibo no post')
add_bullet('Se eles engajarem ou compartilharem, você alcança o público de 440k+ empresas deles de graça')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 9. METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  9. Métricas para Acompanhar (Revisão Mensal)')
doc.add_paragraph()

add_body('Acompanhe pelo LinkedIn Analytics — verifique mensalmente, não diariamente:')

add_table(
    ['Métrica', 'Meta'],
    [
        ['Impressões por post',         '500+ organicamente'],
        ['Taxa de engajamento',         '3%+ (reações + comentários / impressões)'],
        ['Visualizações de perfil após post', 'Pico em 48h'],
        ['Crescimento de seguidores',   '+20–40/mês'],
        ['Envios de Lead Gen Form',     '5–15/mês (com impulsionamento)'],
    ]
)

add_callout('Se um post tiver menos de 1% de engajamento, não o impulsione. Se um post tiver mais de 5%, impulsione imediatamente enquanto está quente.')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 10. BUDGET SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  10. Resumo de Orçamento Mensal')
doc.add_paragraph()

add_table(
    ['Canal', 'Orçamento', 'Objetivo'],
    [
        ['LinkedIn Sponsored (1 post × 7 dias × R$50/dia)', 'R$350',  'Leads diretos (CFO/CEO)'],
        ['Instagram Boost',                                  'R$150',  'Awareness + pool de retargeting'],
        ['Total/mês',                                        'R$500',  ''],
    ]
)

add_callout('A R$500/mês, você precisa de 1 projeto BPO ou BI fechado por trimestre para cobrir um ano de publicidade. Dado o tamanho dos seus contratos, essa matemática é muito favorável.')

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 11. MONTHLY RHYTHM
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  11. Ritmo Mensal de Conteúdo')
doc.add_paragraph()

add_table(
    ['Semana', 'Ação'],
    [
        ['Semana 1', 'Publique conteúdo da semana temática atual (texto + infográfico)'],
        ['Semana 2', 'Publique carrossel + texto da próxima semana temática'],
        ['Semana 3', 'Publique infográfico + carrossel da próxima semana'],
        ['Semana 4', 'Impulsione o melhor post do mês + planeje o próximo tema'],
    ]
)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# 12. ACTION ITEMS
# ══════════════════════════════════════════════════════════════════════════════
add_h1('  12. O que Fazer Agora — Por Prioridade')
doc.add_paragraph()

items = [
    ('1. Salvar o template de segmentação do Lead Gen Form', ' no LinkedIn Campaign Manager — leva 20 minutos, reutilize todo mês'),
    ('2. Agendar os posts da Semana 3 nativamente no LinkedIn:', ' Terça 08h texto, Quarta 12h infográfico, Quinta 08h30 carrossel'),
    ('3. Configurar um boost de R$30/dia', ' no post que tiver melhor desempenho orgânico esta semana'),
    ('4. Criar uma lista de broadcast do WhatsApp Business', ' de clientes atuais e anteriores — comece a compartilhar infográficos imediatamente'),
    ('5. Marcar @Nibo', ' no próximo post BPO'),
    ('6. Subir os carrosséis no Instagram', ' após a publicação no LinkedIn (versão IG: nova legenda curta + 10 hashtags)'),
]

for bold, rest in items:
    add_bullet(rest, bold_prefix=bold)

doc.add_paragraph()
add_callout('O conteúdo já está construído. Esta estratégia transforma tudo em um sistema.')

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r'c:\Users\stern\Documents\myfirstGit\p1p_content_generator\P1P_Estrategia_Publicacao.docx'
doc.save(out)
print(f'Saved: {out}')
